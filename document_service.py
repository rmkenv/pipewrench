import os
import hashlib
from typing import List, Optional, BinaryIO
from fastapi import UploadFile, HTTPException
from pathlib import Path
import PyPDF2
from docx import Document as DocxDocument
import aiofiles
from sqlalchemy.orm import Session
from models.database import Document, User
from config.settings import settings

class DocumentProcessor:
    def __init__(self):
        self.upload_dir = Path(settings.upload_dir)
        self.upload_dir.mkdir(exist_ok=True)

    async def save_uploaded_file(self, file: UploadFile) -> str:
        """Save uploaded file and return file path"""
        # Validate file
        self.validate_file(file)

        # Generate unique filename
        file_hash = hashlib.md5(file.filename.encode()).hexdigest()[:8]
        file_extension = Path(file.filename).suffix
        unique_filename = f"{file_hash}_{file.filename}"
        file_path = self.upload_dir / unique_filename

        # Save file
        async with aiofiles.open(file_path, 'wb') as f:
            content = await file.read()
            await f.write(content)

        return str(file_path)

    def validate_file(self, file: UploadFile):
        """Validate file type and size"""
        if not file.filename:
            raise HTTPException(status_code=400, detail="File name is required")

        file_extension = Path(file.filename).suffix.lower()
        if file_extension not in settings.allowed_extensions:
            raise HTTPException(
                status_code=400, 
                detail=f"File type {file_extension} not allowed. Allowed types: {settings.allowed_extensions}"
            )

        if file.size and file.size > settings.max_file_size:
            raise HTTPException(
                status_code=400, 
                detail=f"File too large. Max size: {settings.max_file_size} bytes"
            )

    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text.strip()
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error extracting PDF text: {str(e)}")

    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = DocxDocument(file_path)
            text = []
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)
            return "\n".join(text)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error extracting DOCX text: {str(e)}")

    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read()
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error extracting TXT text: {str(e)}")

    def extract_text(self, file_path: str) -> str:
        """Extract text from file based on extension"""
        file_extension = Path(file_path).suffix.lower()

        if file_extension == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            return self.extract_text_from_docx(file_path)
        elif file_extension == '.txt':
            return self.extract_text_from_txt(file_path)
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {file_extension}")

    async def process_document(
        self, 
        file: UploadFile, 
        document_type: str,
        document_category: str,
        user: User,
        db: Session
    ) -> Document:
        """Process uploaded document and save to database"""
        # Save file
        file_path = await self.save_uploaded_file(file)

        # Extract text content
        content_text = self.extract_text(file_path)

        # Create database record
        document = Document(
            filename=Path(file_path).name,
            original_filename=file.filename,
            file_path=file_path,
            file_type=document_type,
            document_category=document_category,
            file_size=file.size,
            mime_type=file.content_type,
            content_text=content_text,
            uploaded_by_id=user.id,
            processed=True
        )

        db.add(document)
        db.commit()
        db.refresh(document)

        return document

    def get_documents_by_type(self, db: Session, document_type: str) -> List[Document]:
        """Get all documents of a specific type"""
        return db.query(Document).filter(Document.file_type == document_type).all()

    def get_documents_for_role(self, db: Session, role_related_keywords: List[str]) -> List[Document]:
        """Get documents related to a specific job role"""
        # This is a simple implementation - could be enhanced with better text matching
        documents = []
        for keyword in role_related_keywords:
            role_docs = db.query(Document).filter(
                Document.content_text.ilike(f'%{keyword}%')
            ).all()
            documents.extend(role_docs)

        # Remove duplicates
        seen_ids = set()
        unique_documents = []
        for doc in documents:
            if doc.id not in seen_ids:
                unique_documents.append(doc)
                seen_ids.add(doc.id)

        return unique_documents

# Global document processor instance
document_processor = DocumentProcessor()
