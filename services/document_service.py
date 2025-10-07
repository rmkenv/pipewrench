
"""Document processing service."""
import os
import hashlib
from typing import List, Optional
from fastapi import UploadFile, HTTPException
from pathlib import Path
import logging

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

import aiofiles
from sqlalchemy.orm import Session
from models.database import Document, User
from config.settings import settings

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Service for document processing and text extraction."""
    
    def __init__(self):
        """Initialize document processor."""
        self.upload_dir = Path(settings.upload_dir)
        self.upload_dir.mkdir(parents=True, exist_ok=True)
        logger.info(f"Document upload directory: {self.upload_dir}")

    async def save_uploaded_file(self, file: UploadFile) -> str:
        """Save uploaded file and return file path."""
        try:
            # Validate file
            self.validate_file(file)

            # Generate unique filename
            file_hash = hashlib.md5(file.filename.encode()).hexdigest()[:8]
            file_extension = Path(file.filename).suffix
            unique_filename = f"{file_hash}_{file.filename}"
            file_path = self.upload_dir / unique_filename

            # Save file
            async with aiofiles.open(file_path, 'wb') as f:
                content = await file.read()
                await f.write(content)

            logger.info(f"File saved: {unique_filename}")
            return str(file_path)
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Error saving file: {e}")
            raise HTTPException(status_code=500, detail=f"Error saving file: {str(e)}")

    def validate_file(self, file: UploadFile):
        """Validate file type and size."""
        if not file.filename:
            raise HTTPException(status_code=400, detail="File name is required")

        file_extension = Path(file.filename).suffix.lower()
        if file_extension not in settings.allowed_extensions:
            raise HTTPException(
                status_code=400, 
                detail=f"File type {file_extension} not allowed. Allowed types: {', '.join(settings.allowed_extensions)}"
            )

        # Note: file.size might not be available for all upload types
        # Additional size check can be done after reading

    def extract_text(self, file_path: str) -> str:
        """Extract text from file based on extension."""
        try:
            file_extension = Path(file_path).suffix.lower()
            
            if file_extension == '.pdf':
                return self.extract_text_from_pdf(file_path)
            elif file_extension in ['.doc', '.docx']:
                return self.extract_text_from_docx(file_path)
            elif file_extension == '.txt':
                return self.extract_text_from_txt(file_path)
            else:
                logger.warning(f"Unsupported file type for text extraction: {file_extension}")
                return ""
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return ""

    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file."""
        if not PyPDF2:
            logger.error("PyPDF2 not installed")
            raise HTTPException(status_code=500, detail="PDF processing not available")
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                return text.strip()
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            raise HTTPException(status_code=500, detail=f"Error extracting PDF text: {str(e)}")

    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        if not DocxDocument:
            logger.error("python-docx not installed")
            raise HTTPException(status_code=500, detail="DOCX processing not available")
        
        try:
            doc = DocxDocument(file_path)
            text = []
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    text.append(paragraph.text)
            return "\n".join(text)
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            raise HTTPException(status_code=500, detail=f"Error extracting DOCX text: {str(e)}")

    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                logger.error(f"Error reading TXT file: {e}")
                raise HTTPException(status_code=500, detail=f"Error reading TXT file: {str(e)}")
        except Exception as e:
            logger.error(f"Error extracting TXT text: {e}")
            raise HTTPException(status_code=500, detail=f"Error extracting TXT text: {str(e)}")

    async def process_document(
        self,
        file: UploadFile,
        document_type: str,
        user: User,
        db: Session
    ) -> Document:
        """Process uploaded document and save to database."""
        try:
            # Save file
            file_path = await self.save_uploaded_file(file)
            
            # Get file info
            file_size = os.path.getsize(file_path)
            file_extension = Path(file.filename).suffix.lower()
            
            # Extract text
            content_text = self.extract_text(file_path)
            
            # Create database record
            document = Document(
                filename=Path(file_path).name,
                original_filename=file.filename,
                file_path=file_path,
                file_type=file_extension,
                file_size=file_size,
                content_text=content_text,
                document_type=document_type,
                uploaded_by=user.id
            )
            
            db.add(document)
            db.commit()
            db.refresh(document)
            
            logger.info(f"Document processed: {document.id} - {document.original_filename}")
            return document
            
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Error processing document: {e}")
            db.rollback()
            raise HTTPException(status_code=500, detail=f"Error processing document: {str(e)}")


# Global instance
document_processor = DocumentProcessor()
